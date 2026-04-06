"""
Generate the improved Coastline Digital Solutions Business Plan as a .docx file.
Incorporates all fixes from the review.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 70, 127)  # dark teal-blue
    hs.font.name = 'Calibri'

# ── Helper functions ────────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacing


def bp(text, bold=False):
    """Add a body paragraph."""
    p = doc.add_paragraph(text)
    if bold:
        for r in p.runs:
            r.bold = True
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p


# ════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Business Plan\n')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 70, 127)
run2 = title.add_run('Coastline Digital Solutions')
run2.bold = True
run2.font.size = Pt(22)
run2.font.color.rgb = RGBColor(0, 70, 127)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: April 2026\n').font.size = Pt(12)
info.add_run('Location: British Columbia, Canada (remote-first, serving Canada + international)\n').font.size = Pt(12)
info.add_run('Legal Structure: Sole Proprietorship\n').font.size = Pt(12)
info.add_run('\nVersion 2.0 – Revised').font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Company Overview',
    '3. Services & Sub-Services (Four Pillars)',
    '4. Market Opportunity',
    '5. Business Model & Revenue',
    '6. Target Segment Strategy',
    '7. Operations Plan',
    '8. Marketing & Sales Plan',
    '9. Financial Plan',
    '10. Risk Analysis & Mitigation',
    '11. Milestones & Roadmap',
    '12. KPI Dashboard',
    '13. Long-Term Vision & Exit Strategy',
    '14. Appendices',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

bp('Coastline Digital Solutions helps businesses turn unsolved market frustrations into software, '
   'process changes, or training that makes them the default choice in their niche. Unlike generic '
   'agencies that build whatever clients ask for, we start every engagement with a Market Gap Audit – '
   'a data-driven analysis of customer complaints, competitor blind spots, and regulatory shifts. '
   'Only then do we build a solution, and we build only the #1 gap.')

doc.add_heading('Four Integrated Service Pillars', level=2)
bullet('Business Consultancy – gap-closing operations design')
bullet('Market Analysis – frustration mining & competitor gaps')
bullet('Business Training – teaching various business courses tailored to customer needs (launches Month 9)')
bullet('Digital Solutions – software/apps/websites that fix the #1 gap')

doc.add_heading('Primary Target Segments', level=2)
bullet('Startups & SMEs (0–50 employees) – primary focus in Year 1')
bullet('Corporations (250+ employees) – opportunistic in Year 1, systematic in Year 2')

doc.add_heading('Geographic Scope', level=2)
bullet('Year 1: Canada (all provinces, starting from BC)')
bullet('Year 2: US (East & West Coast)')
bullet('Year 3+: UK, Australia, NZ, EU')

doc.add_heading('Revenue Model', level=2)
bullet('Project fees + monthly retainers (competitive intelligence & gap monitoring)')
bullet('Optional equity/revenue share for high-potential startups (capped at 2–3 deals/year)')

doc.add_heading('Key Financial Targets (Year 1)', level=2)
add_table(doc,
    ['Metric', 'Conservative', 'Realistic', 'Optimistic'],
    [
        ['Revenue (CAD)', '$180k–$250k', '$250k–$350k', '$380k–$450k'],
        ['Gross Margin', '60–70%', '60–70%', '65–75%'],
        ['Breakeven', 'Month 8–10', 'Month 6–8', 'Month 4–6'],
        ['Audit → Build Conversion', '>40%', '>60%', '>70%'],
        ['Net Profit (before founder salary)', '$30k–$60k', '$94k–$150k', '$200k+'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  2. COMPANY OVERVIEW
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Company Overview', level=1)

doc.add_heading('2.1 Legal Name & Registration', level=2)
bp('Coastline Digital Solutions (Sole Proprietorship, Canada). Planned US LLC via Stripe Atlas for USD contracts.')

doc.add_heading('2.2 Mission', level=2)
bp('Turn unsolved market frustrations into software that makes our clients the default choice in their niche – '
   'before their competitors even see the gap.')

doc.add_heading('2.3 Unique Value Proposition', level=2)
p = doc.add_paragraph()
run = p.add_run('"Most software agencies ask what you want. We ask what your market hates – '
                'then build the solution that makes you the only one who fixes it."')
run.italic = True
run.font.size = Pt(12)

doc.add_heading('2.4 Core Differentiators (The Moat)', level=2)
bullet('Paid Market Gap Audit with a "no gap, no fee" guarantee.')
bullet('Build only the #1 gap – no feature bloat, no guessing.')
bullet('Outcome-linked pricing (e.g., "reduce support tickets about X by 70%").')
bullet('Four integrated pillars – consultancy, analysis, training, software – that feed each other.')
bullet('Proprietary methodology for scraping, clustering, and prioritising customer frustrations.')

doc.add_heading('2.5 Defensibility Strategy (How the Moat Deepens Over Time)', level=2)
bp('The methodology itself (scraping reviews, clustering with NLP) is reproducible. '
   'The real moat is built through three compounding assets:')
bullet('Published thought leadership – becoming the recognised authority on "frustration mining" through case studies, blog posts, and speaking engagements.')
bullet('Proprietary sector datasets – each completed audit adds to an internal database of complaints, gaps, and solutions per sector. After 20+ audits, our pattern recognition is unmatched.')
bullet('Internal tooling – custom-built automation that reduces audit delivery time from 7 days to 3 days. Competitors starting from scratch cannot match this speed.')
bullet('Client outcome track record – every successful build creates a referenceable case study with measurable ROI. This compounds trust over time.')

doc.add_heading('2.6 Founder & Team (Initial)', level=2)
add_table(doc,
    ['Role', 'Status'],
    [
        ['Founder (CEO): market analysis, sales, audit delivery', 'Full-time'],
        ['Full-stack developer', '1 contractor (40h/week)'],
        ['Data analyst (frustration mining)', '1 contractor (20h/week, starting Month 3)'],
        ['Legal/accounting', 'On demand'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  3. SERVICES & SUB-SERVICES
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Services & Sub-Services (Four Pillars)', level=1)

# 3.1 Business Consultancy
doc.add_heading('3.1 Business Consultancy', level=2)
bp('Gap-closing operational redesign without code.')
add_table(doc,
    ['Sub-service', 'Problem Solved', 'Deliverable'],
    [
        ['Customer friction audit', 'Support team answers same questions daily', 'Top 5 friction points + process changes to eliminate 3'],
        ['Competitor response playbook', 'Losing deals because competitor has feature Y', 'Playbook to neutralise that advantage'],
        ['Regulatory gap analysis', 'New regulation ignored by competitors', 'Timeline + action plan to turn compliance into marketing'],
        ['Vendor consolidation roadmap', '7 tools that don\'t talk – customers feel delays', 'Integration map + stack to reduce hand-offs by 60%'],
    ])

# 3.2 Market Analysis
doc.add_heading('3.2 Market Analysis', level=2)
bp('Data-proven gaps, not opinions.')
add_table(doc,
    ['Sub-service', 'Problem Solved', 'Deliverable'],
    [
        ['Frustration mining (core)', '"I don\'t know why customers are unhappy"', '500–2,000 complaints clustered, top 3 gaps'],
        ['Competitor feature gap matrix', '"All competitors look the same"', 'Missing features mapped to complaints'],
        ['Price & positioning gap', 'Losing to cheaper competitors', 'Willingness-to-pay analysis + unbundle recommendation'],
        ['Emerging threat scan', 'New startup entered our market', 'Deep dive on weak spots + counter-play'],
    ])

# 3.3 Business Training
doc.add_heading('3.3 Business Training', level=2)
bp('Teach clients to spot gaps themselves. Note: Training pillar launches formally in Month 9, '
   'once 5+ completed audits/builds provide credible curriculum material.')
add_table(doc,
    ['Sub-service', 'Format', 'Fee Range (CAD)'],
    [
        ['Gap hunting workshop (company)', '2-day on-site/remote for up to 10 people', '$8k–$15k'],
        ['Frustration-led requirements (product managers)', '4-week live cohort', '$1.5k/person'],
        ['Competitive moat design (founders)', '1-day executive session', '$5k–$10k'],
        ['Automation opportunity spotting (ops teams)', 'Half-day workshop', '$3k–$5k'],
        ['Individual certification', 'Self-paced + project + certification', '$1k–$2k'],
    ])

# 3.4 Digital Solutions
doc.add_heading('3.4 Digital Solutions', level=2)
bp('Build only the #1 solution to the #1 gap.')
add_table(doc,
    ['Sub-service', 'Description', 'Fee Range (CAD)'],
    [
        ['Unfair Advantage Feature', 'Single high-impact feature solving top frustration', '$25k–$100k'],
        ['Process automation tool', 'Lightweight internal web app for 1 workflow', '$15k–$40k'],
        ['Customer-facing gap closer', 'Portal that solves a specific complaint (e.g., real-time tracker)', '$30k–$80k'],
        ['Clone-proof MVP', 'MVP with deliberate complexity to slow copycats', '$40k–$120k'],
        ['Gap-focused website', 'Website built around a customer problem (e.g., compliance checker)', '$10k–$30k'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  4. MARKET OPPORTUNITY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Market Opportunity', level=1)

doc.add_heading('4.1 Target Sectors (by Frustration Density)', level=2)
add_table(doc,
    ['Sector', 'Evidence of Unsolved Complaints', 'Geography'],
    [
        ['Strata / condo management', 'Owners: "no transparency," "maintenance ignored"', 'Canada, US, Australia'],
        ['Small construction contractors', 'Homeowners: "change orders lost," "no updates"', 'All English-speaking markets'],
        ['Property management (rental)', 'Tenants: "maintenance takes weeks"', 'Canada, US, UK'],
        ['Non-profit grant reporting', 'Grant officers: "inconsistent data," "late reports"', 'Canada, US, EU'],
        ['Cannabis retail', 'Customers: "always out of stock"', 'Canada, US (state-by-state), Germany'],
    ])

doc.add_heading('4.2 Geographic Rollout (Revised – Focused)', level=2)
bp('Year 1 is intentionally focused on Canada to build case studies and refine the methodology before expanding internationally.')
add_table(doc,
    ['Phase', 'Region', 'Timeline'],
    [
        ['1', 'British Columbia (home base)', 'Months 1–6'],
        ['2', 'Ontario, Alberta, Quebec (French-ready)', 'Months 7–12'],
        ['3', 'US (East & West Coast) via US LLC', 'Year 2 (Months 13–24)'],
        ['4', 'UK, Australia, NZ', 'Year 3'],
        ['5', 'EU (German, French localisation)', 'Year 3+'],
    ])

doc.add_heading('4.3 Competitive Landscape', level=2)
add_table(doc,
    ['Competitor Type', 'Their Weakness', 'Our Advantage'],
    [
        ['Generic dev agencies', 'Build what client asks, no market insight', 'We start with data, not requests'],
        ['Freelance consultants', 'No software delivery capability', 'We build the solution end-to-end'],
        ['Large consultancies (Deloitte, etc.)', 'Too expensive for SMEs, slow delivery', 'Agile, fixed-price, SME-friendly'],
        ['Market research firms', 'Reports only, no execution', 'We deliver software + training'],
        ['No-code agencies', 'Limited to simple apps', 'Full-stack, complex automations'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  5. BUSINESS MODEL & REVENUE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Business Model & Revenue', level=1)

doc.add_heading('5.1 Revenue Streams', level=2)
add_table(doc,
    ['Stream', 'Pillar', 'Pricing (CAD)', 'Recurring?'],
    [
        ['Market Gap Audit (full)', 'Market Analysis', '$5k (SME), $15k–$25k (corp)', 'No'],
        ['Gap Snapshot (light)', 'Market Analysis', '$1.5k', 'No'],
        ['Unfair Advantage Build', 'Digital Solutions', '$25k–$100k (SME), $80k–$250k (corp)', 'No'],
        ['Automation tool', 'Digital Solutions', '$15k–$40k (SME), $50k–$150k (corp)', 'No'],
        ['Gap website', 'Digital Solutions', '$10k–$20k', 'No'],
        ['Consulting (friction audit, etc.)', 'Consultancy', '$3k–$12k', 'No'],
        ['Training workshop (company)', 'Training', '$8k–$15k (SME), $30k–$60k (corp)', 'No'],
        ['Competitive intelligence retainer', 'Market Analysis', '$2k–$5k/month (SME), $8k–$12k/month (corp)', 'Yes'],
        ['Gap monitoring retainer (post-build)', 'Digital Solutions', '$2k–$5k/month', 'Yes'],
        ['White-label audit license', 'Market Analysis', '$15k setup + 10% royalty', 'Yes (royalty)'],
    ])

doc.add_heading('5.2 Pricing & Packaging – Startups/SMEs', level=2)
add_table(doc,
    ['Package', 'Included', 'Price (CAD)'],
    [
        ['Gap Snapshot', '50 complaints, top 3 frustrations', '$1,500'],
        ['Full Market Gap Audit', '500–2,000 complaints, competitor gap, mockup', '$5,000'],
        ['Friction Fix (consulting)', 'Top 5 friction points + process fixes', '$3,000'],
        ['Startup Unfair Advantage', 'Full Audit + Unfair Advantage Build', '$25k–$40k'],
        ['SME Automation Sprint', 'Vendor consolidation audit + automation tool', '$15k–$30k'],
        ['Gap-Ready Website', 'Light audit + website as a gap tool', '$10k–$20k'],
        ['Gap Hunting Workshop (2 days)', 'Training for up to 10 people', '$8,000'],
    ])

bp('Payment terms (SME):')
bullet('Under $10k: 100% upfront or 50/50')
bullet('$10k–$40k: 30% / 40% / 30% (deposit / milestone / delivery)')
bullet('Over $40k: 20% / 30% / 30% / 20%')

bp('Equity/Revenue share (guarded):')
bullet('Available only for high-potential startups, capped at 2–3 deals per year.')
bullet('Terms: Reduced cash fee + 2–5% equity OR 5–10% revenue share for 12–24 months.')
bullet('Founder must personally approve. A separate equity evaluation checklist will be used.')

doc.add_heading('5.3 Pricing & Packaging – Corporations', level=2)
add_table(doc,
    ['Package', 'Included', 'Price (CAD)'],
    [
        ['Enterprise Market Gap Audit', '2,000+ complaints, sentiment analysis, ROI projections', '$15k–$25k'],
        ['Competitive Intelligence Retainer', 'Quarterly reports + counter-measures', '$8k–$12k/month'],
        ['Corporate Unfair Advantage Build', 'Enterprise-grade software, SSO, SLA', '$80k–$250k'],
        ['Process Automation (Enterprise)', 'Cross-departmental + ERP/CRM integration', '$50k–$150k'],
        ['Enterprise Cohort Training', 'Train up to 50 employees, certification', '$30k–$60k'],
        ['End-to-End Moat Package', 'All four pillars + 12 months retainer', '$250k–$500k'],
    ])
bp('Payment terms (corp): Net-30 invoicing, 25/25/25/25 or 30/30/30/10. P.O. accepted.')

doc.add_heading('5.4 Bundles (Cross-Pillar Discounts)', level=2)
bp('SME Bundles:', bold=True)
bullet('Gap to MVP: Audit + Build – 10% off → $27k–$36k')
bullet('Friction to Fix: Consulting + Automation – 10% off → $16k–$30k')
bullet('Complete Startup Package: Audit + Build + Workshop – 15% off → $35k–$50k')

bp('Corporate Bundles:', bold=True)
bullet('Market Intelligence Suite: Enterprise Audit + 6 months retainer – 15% off → $55k–$70k')
bullet('Digital Transformation Bundle: Process Automation + Unfair Advantage Build – 10% off → $117k–$360k')
bullet('Strategic Moat (annual): End-to-End + Cohort Training + 12 months retainer – 20% off → $300k–$550k')

doc.add_heading('5.5 International Pricing', level=2)
add_table(doc,
    ['Service', 'Canada (CAD)', 'US (USD)', 'UK (GBP)', 'EU (EUR)'],
    [
        ['Full Audit (SME)', '$5,000', '$4,000', '£3,200', '€3,700'],
        ['Unfair Advantage Build (SME)', '$25k–$40k', '$20k–$32k', '£16k–£26k', '€18k–€30k'],
        ['Competitive Retainer (corp, monthly)', '$8k–$12k', '$6.5k–$10k', '£5.2k–£8k', '€6k–€9k'],
    ])
bp('Add 3–5% currency buffer for quotes. Payment in local currency via Stripe/Wise.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  6. TARGET SEGMENT STRATEGY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Target Segment Strategy', level=1)

doc.add_heading('6.1 Startups & SMEs (0–50 employees, <$5M revenue)', level=2)
bullet('Entry point: Gap Snapshot ($1.5k) → upsell to Full Audit ($5k) → Build ($25k+)')
bullet('Sales channel: Founder-led outbound (LinkedIn, cold email with free frustration snippet), content marketing')
bullet('Decision maker: Founder or head of product')
bullet('Sales cycle: 1–4 weeks')

doc.add_heading('6.2 Corporations (250+ employees, >$50M revenue)', level=2)
bullet('Entry point: Enterprise Audit ($15k–$25k) or Competitive Retainer ($8k–$12k/month)')
bullet('Sales channel: Partnerships with management consultants, referrals, speaking at industry events, account-based marketing')
bullet('Decision maker: VP of Strategy, Director of Ops, CTO')
bullet('Sales cycle: 3–6 months')
bp('Note: Corporate prospecting begins in Month 6, with first corporate engagement expected Month 10–14. '
   'Any earlier corporate win is treated as upside, not planned revenue.')

doc.add_heading('6.3 Upgrade Path', level=2)
add_table(doc,
    ['Current Stage', 'Next Engagement', 'Upsell Trigger'],
    [
        ['Seed startup (Gap Snapshot)', 'Full Audit + Build', '"You need data before building"'],
        ['Growth SME (one build)', 'Competitive retainer', '"Competitors will copy you – let\'s monitor"'],
        ['Established SME', 'Enterprise Audit', '"You now have 100+ employees, time for enterprise-grade"'],
        ['Corporate', 'Multi-year strategic partnership', '"We become your external gap intelligence unit"'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  7. OPERATIONS PLAN
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Operations Plan', level=1)

doc.add_heading('7.1 Service Delivery Process (The Coastline Method)', level=2)
add_table(doc,
    ['Step', 'Owner', 'Duration', 'Deliverable'],
    [
        ['1. Discovery call', 'Founder', '30 min', 'Qualify lead, propose Gap Snapshot or Full Audit'],
        ['2. Sign & pay (50% deposit)', 'Admin', '1 day', 'Contract + invoice'],
        ['3. Data collection', 'Data analyst', '3–5 days', 'Raw scrapes (reviews, forums, support logs)'],
        ['4. Frustration clustering', 'Data analyst + AI', '2 days', 'Clustered complaints, heatmap'],
        ['5. Competitor gap analysis', 'Founder', '2 days', 'Gap matrix + prioritisation'],
        ['6. Software concept (if build)', 'Founder + dev', '3 days', 'Mockup, user stories, cost estimate'],
        ['7. Delivery & presentation', 'Founder', '1 day', 'Audit report + concept deck'],
        ['8. Build phase (if sold)', 'Dev + founder', '4–12 weeks', 'Software launch + outcome tracking'],
    ])

doc.add_heading('7.2 Founder Capacity Limits & Hiring Triggers', level=2)
bp('The founder is the bottleneck in Year 1. To avoid over-commitment, the following limits apply:')
add_table(doc,
    ['Scenario', 'Max Concurrent Load', 'Action Required'],
    [
        ['Audits only', '3 simultaneous audits', 'At 4+ audits in pipeline: pause outbound, deliver first'],
        ['Audit + Build', '2 audits + 1 active build', 'If 2nd build closes: contractor dev must own delivery'],
        ['Builds only', '2 active builds max', 'At 3rd build: hire 2nd developer (contractor)'],
        ['Training added', '1 workshop/month max', 'More demand: hire part-time facilitator or delay'],
    ])
bp('Trigger for first full-time hire: sustained 3+ months of >$25k/month revenue.')

doc.add_heading('7.3 Client Retention & Churn Strategy', level=2)
bp('For retainer clients, the following retention playbook applies:')
bullet('Monthly value report: Every retainer client receives a one-page summary showing new gaps found, competitor moves tracked, and actions recommended.')
bullet('Quarterly strategy call: 30-minute review of the competitive landscape with the client\'s decision-maker.')
bullet('Upsell on change: When a significant market shift occurs (new competitor, regulatory change), proactively propose a response — this renews the retainer\'s perceived value.')
bullet('Graduation path: When a client outgrows monitoring, transition them to an annual "Strategic Review" engagement ($15k–$25k) rather than losing them entirely.')
bullet('Target retainer churn rate: <15% annually.')

doc.add_heading('7.4 Technology Stack', level=2)
add_table(doc,
    ['Purpose', 'Tools'],
    [
        ['Scraping', 'BrightData, Apify, custom Python (BeautifulSoup)'],
        ['NLP/Clustering', 'OpenAI API (GPT-4 for classification), spaCy'],
        ['Project management', 'Notion, ClickUp'],
        ['Development', 'React/Node.js, Python (Django/FastAPI), PostgreSQL, AWS/GCP'],
        ['Payments', 'Stripe, Wise, Interac (Canada)'],
        ['CRM', 'Pipedrive or HubSpot (free tier initially)'],
        ['Accounting', 'Xero + Hubdoc'],
    ])

doc.add_heading('7.5 Legal & Compliance', level=2)
bullet('Canada: Federal incorporation, register extra-provincially in ON, AB, QC as needed.')
bullet('US: US LLC via Stripe Atlas or LegalZoom (Delaware or Wyoming).')
bullet('Contracts: Custom SOW for each engagement. IP assignment: client owns software, we own methodology.')
bullet('Data privacy: Comply with PIPEDA (Canada), state privacy laws (US), GDPR (EU clients – add DPA).')
bullet('Insurance: Professional liability ($1M), cyber liability.')

doc.add_heading('7.6 Staffing Plan (Years 1–2)', level=2)
add_table(doc,
    ['Role', 'Year 1', 'Year 2'],
    [
        ['Founder (CEO, sales, audit lead)', '1', '1'],
        ['Full-stack developer', '1 (contractor)', '2 (1 employee + 1 contractor)'],
        ['Data analyst (part-time)', '1 (contractor, 20h/week, from Month 3)', '1 (full-time)'],
        ['Project manager', '0', '1'],
        ['Sales development (outbound)', '0', '1'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  8. MARKETING & SALES PLAN
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Marketing & Sales Plan', level=1)

doc.add_heading('8.1 Lead Generation', level=2)
add_table(doc,
    ['Tactic', 'Description', 'Cost (Monthly)'],
    [
        ['Free Frustration Snapshot', 'Publish "Top 5 things [sector] customers hate" – no pitch, just data. Share on LinkedIn, industry forums.', '$0 (time)'],
        ['Outbound (SME)', 'Find businesses with bad reviews. Send: "I analysed your 47 negative reviews. The #1 complaint is X. Want the 2-page summary?"', '$500 (LinkedIn Sales Navigator)'],
        ['Content engine', 'Case studies of audits → builds (anonymised). Blog posts: "How we found a $2M gap in [sector]."', '$0 (founder writes)'],
        ['Partnerships (corp)', 'Partner with management consultants who don\'t build software. 15% referral commission.', '15% of deal'],
        ['Speaking & webinars', '"Market Gap Hunting 101" webinar for startup accelerators and local incubators.', '$0'],
        ['Paid channels (from Month 4)', 'LinkedIn ads, Google Ads targeting "competitor analysis service," sponsor niche industry newsletters.', '$1k–$2k/month'],
    ])

doc.add_heading('8.2 Sales Process', level=2)
bp('For SMEs (founder-led):', bold=True)
bullet('Connect on LinkedIn with a personalised note referencing a specific complaint.')
bullet('Offer free 15-min "frustration spot check" (analyse 3 reviews live).')
bullet('Propose Gap Snapshot ($1.5k) or Full Audit ($5k).')
bullet('Deliver audit, then propose build.')

bp('For Corporations (longer cycle):', bold=True)
bullet('Warm intro via partner consultant.')
bullet('Offer paid pilot: Enterprise Audit for $15k (discounted from $25k).')
bullet('Present audit to VP level.')
bullet('Propose retainer + build as a 6-month program.')

doc.add_heading('8.3 Conversion Metrics (Targets)', level=2)
add_table(doc,
    ['Metric', 'Target'],
    [
        ['Gap Snapshot → Full Audit', '>40%'],
        ['Full Audit → Build (SME)', '>60%'],
        ['Full Audit → Retainer (corp)', '>50%'],
        ['Average build size (SME)', '$35k'],
        ['Average annual retainer (corp)', '$100k'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  9. FINANCIAL PLAN
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Financial Plan', level=1)

doc.add_heading('9.1 Startup Costs (CAD)', level=2)
add_table(doc,
    ['Expense', 'Amount'],
    [
        ['Incorporation (Federal)', '$500'],
        ['US LLC registration', '$1,000'],
        ['Legal (contract templates, DPAs)', '$2,000'],
        ['Laptop, software licenses', '$3,000'],
        ['Initial marketing (LinkedIn ads test)', '$1,000'],
        ['Professional liability insurance (1 year)', '$2,500'],
        ['Total', '$10,000'],
    ])

doc.add_heading('9.2 Cash Flow & Runway Analysis', level=2)
bp('Critical assumption: Revenue will not arrive on Day 1. The founder must have sufficient personal runway '
   'to cover both business burn and personal expenses during the ramp-up period.')
add_table(doc,
    ['Item', 'Monthly Cost (CAD)', 'Months Needed', 'Total'],
    [
        ['Business operating costs', '$13,000', '4–6 months', '$52k–$78k'],
        ['Founder personal expenses (estimated)', '$4,000–$6,000', '4–6 months', '$16k–$36k'],
        ['Startup costs (one-time)', '—', '—', '$10,000'],
        ['Total runway required', '', '', '$78k–$124k'],
    ])
bp('Funding source: Personal savings and/or line of credit. No external funding planned for Year 1.')
bp('Recommendation: Have a minimum of $80k CAD in accessible funds before launching. '
   'This provides 6 months of runway with zero revenue — enough to reach breakeven under all scenarios.')

doc.add_heading('9.3 Monthly Operating Costs (Lean, Year 1)', level=2)
add_table(doc,
    ['Category', 'Cost (CAD)'],
    [
        ['Developer (contractor, 40h/week)', '$8,000'],
        ['Data analyst (contractor, 20h/week, from Month 3)', '$3,000'],
        ['Cloud hosting (AWS, per project avg)', '$500'],
        ['Scraping tools (BrightData, Apify)', '$400'],
        ['OpenAI API (NLP)', '$200'],
        ['CRM & project management', '$100'],
        ['Accounting (Xero + Hubdoc)', '$60'],
        ['Marketing (content boost, LinkedIn, paid ads from Month 4)', '$500–$2,000'],
        ['Legal/accounting (amortised)', '$200'],
        ['Total monthly burn', '~$13,000–$14,500'],
    ])

doc.add_heading('9.4 Revenue Forecast – Year 1', level=2)

bp('Scenario A: Conservative (slow start, lower conversion)', bold=True)
bp('Assumptions: First client Month 3. Audit→Build conversion 35%. No corporate revenue. No training revenue.')
add_table(doc,
    ['Month', 'Activity', 'Revenue (CAD)', 'Cumulative'],
    [
        ['1–2', 'Setup, outbound begins, no revenue', '$0', '$0'],
        ['3', '2 Gap Snapshots ($1.5k each)', '$3k', '$3k'],
        ['4', '1 Full Audit ($5k) + 1 Gap Snapshot ($1.5k)', '$6.5k', '$9.5k'],
        ['5', '1 Full Audit ($5k) + 1 Build deposit (30% of $25k)', '$12.5k', '$22k'],
        ['6', 'Build milestone ($10k) + 1 Audit ($5k)', '$15k', '$37k'],
        ['7', 'Build final ($7.5k) + 1 Audit ($5k) + retainer ($2k)', '$14.5k', '$51.5k'],
        ['8', '2 Audits ($10k) + 1 new Build deposit ($7.5k) + retainer ($2k)', '$19.5k', '$71k'],
        ['9', 'Build milestone ($10k) + 1 Audit ($5k) + retainer ($2k)', '$17k', '$88k'],
        ['10', 'Build final ($7.5k) + 2 Audits ($10k) + retainer ($2k)', '$19.5k', '$107.5k'],
        ['11', '2 Audits ($10k) + 1 Build deposit ($9k) + retainer ($2k)', '$21k', '$128.5k'],
        ['12', 'Build milestones ($12k) + 2 Audits ($10k) + retainer ($2k)', '$24k', '$152.5k'],
    ])
bp('Conservative Year 1 Total: ~$150k–$180k CAD')
bp('Conservative Year 1 Net: ~$0–$25k (before founder salary) — essentially breakeven.')

bp('')
bp('Scenario B: Realistic (steady ramp, 60% conversion)', bold=True)
add_table(doc,
    ['Month', 'Activity', 'Revenue (CAD)', 'Cumulative'],
    [
        ['1', 'Setup, no revenue', '$0', '$0'],
        ['2', '2 Gap Snapshots + 1 Full Audit', '$8k', '$8k'],
        ['3', '1 Full Audit + 1 Build deposit (30% of $30k)', '$14k', '$22k'],
        ['4', 'Build milestone ($12k) + 1 Audit ($5k)', '$17k', '$39k'],
        ['5', 'Build final ($9k) + 1 Audit ($5k) + 1 Gap Snapshot ($1.5k)', '$15.5k', '$54.5k'],
        ['6', '2 Audits ($10k) + 1 new Build deposit ($9k) + retainer ($2k)', '$21k', '$75.5k'],
        ['7', 'Build milestone & final ($21k) + 1 Audit ($5k) + retainer ($2k)', '$28k', '$103.5k'],
        ['8', '2 Audits ($10k) + 1 Build deposit ($12k) + retainer ($2k)', '$24k', '$127.5k'],
        ['9', 'Build milestone ($16k) + 1 Audit + workshop ($13k) + retainer ($2k)', '$31k', '$158.5k'],
        ['10', 'Build final ($12k) + 2 Audits ($10k) + new Build deposit ($9k) + retainer ($2k)', '$33k', '$191.5k'],
        ['11', '1 Enterprise Audit ($15k) + SME Build final ($9k) + retainer ($2k)', '$26k', '$217.5k'],
        ['12', '2 Audits ($10k) + 1 Corp Build deposit ($15k) + retainer ($8k)', '$33k', '$250.5k'],
    ])
bp('Realistic Year 1 Total: ~$250k CAD')
bp('Realistic Year 1 Net: ~$94k (before founder salary)')

bp('')
bp('Scenario C: Optimistic (fast sales, early corporate win)', bold=True)
bp('Year 1 Total: $380k–$450k. Net profit: $200k+.')
bp('This scenario requires closing a corporate retainer by Month 8 and 2+ concurrent builds from Month 5 onward. Achievable but requires either strong network or lucky timing.')

doc.add_heading('9.5 Year 2–3 Projection', level=2)
add_table(doc,
    ['Year', 'Revenue (CAD)', 'Primary Drivers'],
    [
        ['1', '$180k–$400k', 'SME audits & builds, first retainers'],
        ['2', '$600k–$1M', 'Repeat SME clients, 3–5 corporate retainers, US entry'],
        ['3', '$1.5M–$2.5M', 'Corporate accounts, white-label licensing, international'],
    ])

doc.add_heading('9.6 Breakeven Analysis', level=2)
bullet('Monthly fixed costs (after revenue starts): ~$13k')
bullet('Average gross margin per project: 65% (labour & hosting)')
bullet('Breakeven monthly revenue needed: $13k / 0.65 ≈ $20k/month')
bullet('Breakeven point: Month 6–10 depending on scenario')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  10. RISK ANALYSIS & MITIGATION
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Risk Analysis & Mitigation', level=1)

add_table(doc,
    ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
    [
        ['Low audit → build conversion', 'Medium', 'High',
         'Offer "audit fee credited to build"; improve audit quality with better data; track conversion weekly'],
        ['Long corporate sales cycles', 'High', 'Medium',
         'Build pipeline of 10+ prospects; use partnerships; do not depend on corporate revenue in Year 1'],
        ['Developer turnover', 'Medium', 'Medium',
         'Contract with 2 part-time devs; document code; offer equity to key dev'],
        ['Founder burnout / bottleneck', 'High', 'High',
         'Enforce capacity limits (Section 7.2); hire when triggers are met; take 1 week off per quarter'],
        ['Competitor copies methodology', 'Medium', 'Low',
         'Build moat through published content, proprietary datasets, and speed (Section 2.5)'],
        ['Revenue comes slower than forecast', 'Medium', 'High',
         'Maintain 6 months runway; have "lean mode" budget ($8k/month) ready if needed'],
        ['Currency fluctuation (international)', 'Low', 'Low',
         'Quote in client\'s currency with 3–5% buffer; use Wise/Stripe automatic conversion'],
        ['Data privacy breach', 'Low', 'High',
         'Use DPAs, host sensitive data only with client\'s cloud, never store PII unnecessarily'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  11. MILESTONES & ROADMAP
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Milestones & Roadmap', level=1)

doc.add_heading('Year 1 (2026–2027)', level=2)
add_table(doc,
    ['Month', 'Milestone'],
    [
        ['1', 'Incorporate federally, set up bank account, Stripe, website (landing page)'],
        ['2', 'Publish first "Frustration Snapshot" (strata sector). Close first Gap Snapshot client.'],
        ['3', 'Deliver first Full Audit. Close first build. Hire part-time data analyst.'],
        ['4', 'Launch first software build. Document case study. Begin testing paid ads ($1k/month).'],
        ['5', 'Register US LLC. Start building US prospect list (WA, CA).'],
        ['6', 'Sign first monthly retainer (gap monitoring). Begin corporate prospecting.'],
        ['7', 'Deliver 2nd software build. 5+ completed audits.'],
        ['8', 'Partner with 2 management consultants (referral agreements).'],
        ['9', 'Launch "Gap Hunting Workshop" as a product. Sell to first 2 SMEs.'],
        ['10–11', 'First corporate Enterprise Audit (discounted pilot).'],
        ['12', 'Year 1 revenue review. Evaluate full-time developer hire. Plan Year 2 US expansion.'],
    ])

doc.add_heading('Year 2–3', level=2)
add_table(doc,
    ['Year', 'Milestones'],
    [
        ['2', 'Expand to Ontario & US (East Coast). Launch self-paced certification course. Hire sales development rep. First corporate retainer ($8k+/month). Target $600k–$1M revenue.'],
        ['3', 'Enter UK via local partner. White-label audit licensing to 5 agencies. Corporate accounts >10. Target $1.5M–$2.5M revenue.'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  12. KPI DASHBOARD (NEW)
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('12. KPI Dashboard', level=1)

bp('Track these metrics weekly (W), monthly (M), or quarterly (Q) from Day 1.')

add_table(doc,
    ['KPI', 'Frequency', 'Target (Year 1)', 'Why It Matters'],
    [
        ['Pipeline value (total proposals out)', 'W', '>3x monthly revenue target', 'Predicts future revenue'],
        ['Leads generated (inbound + outbound)', 'W', '10–15/week by Month 3', 'Measures marketing effectiveness'],
        ['Discovery calls booked', 'W', '3–5/week by Month 3', 'Sales activity indicator'],
        ['Gap Snapshot → Full Audit conversion', 'M', '>40%', 'Validates entry offer'],
        ['Full Audit → Build conversion', 'M', '>60%', 'THE key metric — proves methodology works'],
        ['Average deal size (SME builds)', 'M', '$35k', 'Revenue efficiency'],
        ['Time to close (days from first contact to signed contract)', 'M', '<30 days (SME)', 'Sales cycle health'],
        ['Monthly recurring revenue (retainers)', 'M', '$4k+ by Month 8', 'Stability and predictability'],
        ['Client NPS / satisfaction score', 'Q', '>8/10', 'Referral likelihood'],
        ['Retainer churn rate', 'Q', '<15% annually', 'Revenue retention'],
        ['Founder utilisation (% of time on billable work)', 'M', '50–70%', 'Prevents burnout if >80%'],
        ['Cash in bank (months of runway remaining)', 'M', '>3 months always', 'Survival metric'],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  13. LONG-TERM VISION & EXIT STRATEGY (NEW)
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('13. Long-Term Vision & Exit Strategy', level=1)

doc.add_heading('13.1 Vision (3–5 Years)', level=2)
bp('Coastline Digital Solutions aims to become the leading "frustration-to-software" firm for SMEs in '
   'English-speaking markets. By Year 3, the company should have:')
bullet('50+ completed audits and 20+ software builds across 5+ sectors')
bullet('A team of 8–12 people (developers, analysts, sales)')
bullet('$2M+ annual revenue with 30%+ net margin')
bullet('A recognisable brand in the "market gap" niche through published content and speaking')

doc.add_heading('13.2 Strategic Options at Year 3–5', level=2)
add_table(doc,
    ['Option', 'Description', 'Conditions'],
    [
        ['Lifestyle business', 'Founder takes $200k–$400k/year in profit. Team stays small (5–10). Focus on quality over growth.', 'If revenue plateaus at $1M–$2M and founder prefers lifestyle.'],
        ['Growth play', 'Raise seed/Series A funding. Build proprietary SaaS product (e.g., "Frustration Mining Platform"). Scale to $10M+ ARR.', 'If a clear SaaS opportunity emerges from repeat audit patterns.'],
        ['Acquisition target', 'Position for acquisition by a larger consultancy or software firm wanting a "gap methodology" capability.', 'If 3+ corporate clients and a strong brand are established.'],
        ['Franchise / license model', 'License the Coastline methodology and tools to independent consultants globally for a fee + royalty.', 'If white-label licensing proves successful in Year 2–3.'],
    ])

bp('Decision point: Revisit this section at the end of Year 2 with real data. The choice depends on '
   'revenue trajectory, founder preferences, and whether a SaaS product opportunity has emerged.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  14. APPENDICES
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('14. Appendices (To Be Created)', level=1)
bullet('Appendix A: Sample Market Gap Audit report (redacted)')
bullet('Appendix B: Service contracts & SOW templates')
bullet('Appendix C: Founder bio & key team resumes')
bullet('Appendix D: Competitor matrix (detailed)')
bullet('Appendix E: Technology stack diagram')
bullet('Appendix F: Case study – StrataKey (first internal product)')
bullet('Appendix G: Equity/revenue share evaluation checklist')
bullet('Appendix H: "Lean mode" contingency budget ($8k/month)')

doc.add_paragraph()
doc.add_paragraph()

# ── FINAL NOTE ──
doc.add_heading('Final Note', level=1)
bp('This business plan is a living document. As you complete each audit and build, update the '
   'financial forecasts with real data. The most important metric to track from Day 1 is '
   'audit → build conversion rate – that tells you if your gap-finding methodology is working.')

bp('Next immediate actions (next 7 days):', bold=True)
bullet('Incorporate Coastline Digital Solutions Inc. (federal).')
bullet('Set up a simple landing page with the UVP and "Request a Gap Snapshot" button.')
bullet('Publish the first free "Frustration Snapshot" for BC strata managers on LinkedIn.')
bullet('Reach out to 5 local SMEs with negative reviews offering a free 3-review analysis.')
bullet('Confirm personal runway: verify $80k+ CAD is accessible before committing to full launch.')

bp('Once you complete these, you\'ll have your first leads.')

# ── Save ────────────────────────────────────────────────────────────────
output_path = r'e:\Coastline Digital Solutions\Coastline_Digital_Solutions_Business_Plan_v2.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
